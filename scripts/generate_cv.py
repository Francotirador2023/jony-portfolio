import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether, ListFlowable, ListItem
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

def generate_docx():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
    # Styles
    primary_color = RGBColor(14, 116, 144) # Teal/Cyan #0e7490
    dark_color = RGBColor(15, 23, 42)     # Slate 900
    gray_color = RGBColor(71, 85, 105)    # Slate 600
    
    # Title / Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run("Jony Rolando Franco Ayte")
    r_name.font.name = "Arial"
    r_name.font.size = Pt(18)
    r_name.font.bold = True
    r_name.font.color.rgb = dark_color
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_contact = p_contact.add_run("Lima, Perú  |  +51 928 926 775  |  francojonysenati@gmail.com\nlinkedin.com/in/jony-franco-163521358  |  Portafolio: jony-portfolio.vercel.app")
    r_contact.font.name = "Arial"
    r_contact.font.size = Pt(9.5)
    r_contact.font.color.rgb = gray_color
    
    def add_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = primary_color
        
    def add_subheading(left_text, right_text="", is_bold=True, is_italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        r_left = p.add_run(left_text)
        r_left.font.name = "Arial"
        r_left.font.size = Pt(10)
        r_left.font.bold = is_bold
        r_left.font.italic = is_italic
        r_left.font.color.rgb = dark_color
        
        if right_text:
            r_right = p.add_run(f"  ({right_text})")
            r_right.font.name = "Arial"
            r_right.font.size = Pt(9.5)
            r_right.font.italic = True
            r_right.font.color.rgb = gray_color
            
    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        if bold_prefix:
            r_prefix = p.add_run(bold_prefix)
            r_prefix.font.name = "Arial"
            r_prefix.font.size = Pt(9.5)
            r_prefix.font.bold = True
            r_prefix.font.color.rgb = dark_color
        r_text = p.add_run(text)
        r_text.font.name = "Arial"
        r_text.font.size = Pt(9.5)
        r_text.font.color.rgb = dark_color

    # 1. PERFIL PROFESIONAL
    add_heading("PERFIL PROFESIONAL")
    p_profile = doc.add_paragraph()
    p_profile.paragraph_format.space_before = Pt(2)
    p_profile.paragraph_format.space_after = Pt(6)
    r_prof = p_profile.add_run(
        "Data Engineer y Desarrollador de Software especializado en la arquitectura, integración y automatización de pipelines de datos escalables, diseño de APIs y microservicios. Experto Key User en SAP MM (Logística y Materiales), con sólida experiencia en el diseño de soluciones de integración con APIs (REST), procesamiento de datos estructurados y no estructurados, e implementación de sistemas basados en la nube (AWS, GCP, Azure). Orientado a optimizar la disponibilidad y la calidad de los datos para su consumo en analítica avanzada, modelos de Machine Learning y toma de decisiones estratégicas."
    )
    r_prof.font.name = "Arial"
    r_prof.font.size = Pt(9.5)
    r_prof.font.color.rgb = dark_color

    # 2. EXPERIENCIA LABORAL
    add_heading("EXPERIENCIA LABORAL")
    
    # Ccapac Sistemas
    add_subheading("Ccapac Sistemas — Lima, Perú")
    add_subheading("Asesor en Desarrollo de APIs", "2026 – Presente", is_bold=False, is_italic=True)
    add_bullet(" Asesoramiento técnico y consultoría en diseño, arquitectura y desarrollo de APIs REST y microservicios para optimizar la interoperabilidad de sistemas.", bold_prefix="Arquitectura de APIs: ")
    add_bullet(" Desarrollo e integración de componentes backend en PHP y consumo desde interfaces en Angular.", bold_prefix="Desarrollo Backend & Frontend: ")
    add_bullet(" Implementación de buenas prácticas de desarrollo, seguridad y optimización en servicios web.", bold_prefix="Buenas Prácticas & Calidad: ")
    
    # Toulouse Lautrec
    add_subheading("Toulouse Lautrec — Lima, Perú")
    add_subheading("Ingeniero de Datos / Analista", "Octubre 2024 – Junio 2025", is_bold=False, is_italic=True)
    add_bullet(" Diseñé y mantuve pipelines de Extracción, Transformación y Carga (ETL) para grandes volúmenes de datos usando Python y SQL, garantizando integridad y disponibilidad para analítica secundaria.", bold_prefix="Pipelines ETL: ")
    add_bullet(" Automaticé reportes y flujos de datos operativos mediante scripts programados en Python, disminuyendo la carga operativa en un 30% mensual.", bold_prefix="Automatización: ")
    add_bullet(" Colaboré en la integración de datos centralizados y diseño de dashboards interactivos de seguimiento de KPIs en Power BI y Tableau.", bold_prefix="Visualización & BI: ")

    # Proyectos Independientes
    add_subheading("Proyectos de Data Engineering & Automatización", "Independiente / Freelance")
    add_subheading("Data Engineer", "2024 – Presente", is_bold=False, is_italic=True)
    add_bullet(" Arquitecté y operé flujos de trabajo en n8n/Make integrando la API de Gemini (LLMs) para procesar, limpiar y calificar leads B2B automáticamente, ingestando datos al CRM y acelerando el ciclo en un 40%.", bold_prefix="Pipelines con IA (n8n & LLMs): ")
    add_bullet(" Implementación de bases de datos resilientes en la nube (AWS/GCP/Azure) para alta disponibilidad en aplicaciones en producción (jofastsa.com).", bold_prefix="Infraestructura Cloud: ")
    add_bullet(" Procesamiento distribuido y pruebas de arquitecturas escalables con Apache Spark, Hadoop y contenerización en Docker.", bold_prefix="Big Data: ")

    # 3. EDUCACIÓN
    add_heading("EDUCACIÓN")
    add_subheading("Ing. de Sistemas Computacionales", "2026 – Presente (En curso)")
    add_bullet("Universidad Privada del Norte (UPN) — Lima, Perú")
    
    add_subheading("Técnico en Desarrollo de Sistemas de Información", "2022 – 2025 (Egresado / Culminado)")
    add_bullet("Instituto Sabio Nacional Antúnez de Mayolo — Lima, Perú")

    # 4. CERTIFICACIONES Y ESPECIALIZACIONES
    add_heading("CERTIFICACIONES Y ESPECIALIZACIONES")
    add_bullet("SAP MM - Logística y Materiales (Usuario Experto / Key User) | Global Tecnologías Academy | 2026")
    add_bullet("Ciencia de Datos aplicada a la Logística — Specialty | Ingenium | 2025")
    add_bullet("Data Science y Machine Learning — Specialty | Toulouse Lautrec | 2024 – 2025")

    # 5. HABILIDADES TÉCNICAS
    add_heading("HABILIDADES TÉCNICAS Y HERRAMIENTAS")
    add_bullet("PHP, Angular, API REST, Microservicios, Java, FastAPI.", bold_prefix="APIs & Desarrollo Web: ")
    add_bullet("Data Pipelines, Python, SQL, Apache Spark, Hadoop, n8n/Make, Web Scraping.", bold_prefix="Data Engineering & ETL: ")
    add_bullet("PostgreSQL, MySQL, MongoDB, Data Warehousing.", bold_prefix="Bases de Datos: ")
    add_bullet("AWS, Google Cloud Platform (GCP), Microsoft Azure, Docker, Git, CI/CD.", bold_prefix="Cloud & DevOps: ")
    add_bullet("Power BI, Tableau, Streamlit, Metabase.", bold_prefix="Visualización & BI: ")
    add_bullet("SAP MM (Key User), Gestión de Materiales y Procesos Logísticos.", bold_prefix="ERP & Dominio de Negocio: ")

    doc.save("CV_Jony_Data_Engineer.docx")
    print("CV_Jony_Data_Engineer.docx generated successfully.")

def generate_pdf():
    pdf_filename = "public/cv.pdf"
    
    # Page configuration (Letter, 28pt margins for perfect 1-page fit)
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        leftMargin=28,
        rightMargin=28,
        topMargin=26,
        bottomMargin=24
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Palette
    c_primary = colors.HexColor("#0284c7")     # Ocean primary / Sky 600
    c_dark = colors.HexColor("#0f172a")        # Dark slate 900
    c_sub = colors.HexColor("#334155")         # Slate 700
    c_muted = colors.HexColor("#64748b")       # Slate 500
    c_line = colors.HexColor("#cbd5e1")        # Slate 300
    
    # Typography
    styles.add(ParagraphStyle(
        name='CVName',
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=17,
        textColor=c_dark,
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='CVContact',
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=c_sub,
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='CVSectionHeader',
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=12,
        textColor=c_primary,
        spaceBefore=4,
        spaceAfter=1,
        keepWithNext=True
    ))
    
    styles.add(ParagraphStyle(
        name='CVJobTitle',
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=10.5,
        textColor=c_dark,
        keepWithNext=True
    ))
    
    styles.add(ParagraphStyle(
        name='CVJobMeta',
        fontName='Helvetica-Oblique',
        fontSize=8,
        leading=10.5,
        textColor=c_muted,
        alignment=TA_RIGHT,
        keepWithNext=True
    ))
    
    styles.add(ParagraphStyle(
        name='CVBody',
        fontName='Helvetica',
        fontSize=8,
        leading=10.5,
        textColor=c_sub,
        alignment=TA_JUSTIFY
    ))
    
    styles.add(ParagraphStyle(
        name='CVBullet',
        fontName='Helvetica',
        fontSize=7.8,
        leading=10,
        textColor=c_sub,
        leftIndent=10,
        firstLineIndent=-7,
        spaceAfter=1
    ))
    
    styles.add(ParagraphStyle(
        name='CVSkillLabel',
        fontName='Helvetica-Bold',
        fontSize=7.8,
        leading=10,
        textColor=c_dark
    ))
    
    styles.add(ParagraphStyle(
        name='CVSkillVal',
        fontName='Helvetica',
        fontSize=7.8,
        leading=10,
        textColor=c_sub
    ))

    story = []
    
    # Header
    story.append(Paragraph("<b>JONY ROLANDO FRANCO AYTE</b>", styles['CVName']))
    story.append(Spacer(1, 1))
    contact_text = "Lima, Perú &nbsp;|&nbsp; +51 928 926 775 &nbsp;|&nbsp; <b>francojonysenati@gmail.com</b> &nbsp;|&nbsp; <b>linkedin.com/in/jony-franco-163521358</b> &nbsp;|&nbsp; <b>jony-portfolio.vercel.app</b>"
    story.append(Paragraph(contact_text, styles['CVContact']))
    story.append(HRFlowable(width="100%", thickness=1.2, color=c_primary, spaceBefore=2, spaceAfter=3))
    
    # 1. PERFIL PROFESIONAL
    story.append(Paragraph("PERFIL PROFESIONAL", styles['CVSectionHeader']))
    story.append(HRFlowable(width="100%", thickness=0.4, color=c_line, spaceBefore=1, spaceAfter=2))
    prof_text = (
        "<b>Data Engineer y Desarrollador de Software</b> especializado en la arquitectura, integración y automatización de "
        "pipelines de datos escalables, diseño de APIs y microservicios. Experto <b>Key User en SAP MM</b> (Logística y Materiales), "
        "con sólida experiencia en diseño de integraciones API REST, procesamiento de datos estructurados/no estructurados y despliegue "
        "en la nube (AWS, GCP, Azure). Enfocado en maximizar la calidad y disponibilidad de datos para analítica avanzada, "
        "Machine Learning y optimización estratégica de procesos de negocio."
    )
    story.append(Paragraph(prof_text, styles['CVBody']))
    
    # 2. EXPERIENCIA LABORAL
    story.append(Paragraph("EXPERIENCIA LABORAL", styles['CVSectionHeader']))
    story.append(HRFlowable(width="100%", thickness=0.4, color=c_line, spaceBefore=1, spaceAfter=2))
    
    # Ccapac Sistemas
    t_data1 = [
        [Paragraph("<b>Asesor en Desarrollo de APIs</b> — Ccapac Sistemas", styles['CVJobTitle']),
         Paragraph("2026 – Presente | Lima, Perú", styles['CVJobMeta'])]
    ]
    t1 = Table(t_data1, colWidths=[395, 160])
    t1.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 0), ('TOPPADDING', (0,0), (-1,-1), 0), ('LEFTPADDING', (0,0), (-1,-1), 0)]))
    story.append(t1)
    story.append(Paragraph("• <b>Arquitectura de APIs & Microservicios:</b> Asesoramiento técnico y consultoría en diseño, modelado e integración de APIs REST y microservicios.", styles['CVBullet']))
    story.append(Paragraph("• <b>Desarrollo Backend & Frontend:</b> Construcción e integración de servicios en PHP y consumo dinámico desde interfaces modernas con Angular.", styles['CVBullet']))
    story.append(Paragraph("• <b>Optimización & Buenas Prácticas:</b> Estandarización de arquitectura de software, seguridad de endpoints y buenas prácticas de rendimiento.", styles['CVBullet']))
    
    # Toulouse Lautrec
    t_data2 = [
        [Paragraph("<b>Ingeniero de Datos / Analista</b> — Toulouse Lautrec", styles['CVJobTitle']),
         Paragraph("Oct 2024 – Jun 2025 | Lima, Perú", styles['CVJobMeta'])]
    ]
    t2 = Table(t_data2, colWidths=[395, 160])
    t2.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 0), ('TOPPADDING', (0,0), (-1,-1), 0), ('LEFTPADDING', (0,0), (-1,-1), 0)]))
    story.append(t2)
    story.append(Paragraph("• <b>Pipelines ETL:</b> Diseñé y mantuve pipelines ETL para grandes volúmenes de datos con Python y SQL, garantizando integridad para analítica.", styles['CVBullet']))
    story.append(Paragraph("• <b>Automatización de Procesos:</b> Creé scripts en Python para automatizar flujos operativos y reportes periódicos, reduciendo la carga manual un <b>30% mensual</b>.", styles['CVBullet']))
    story.append(Paragraph("• <b>Business Intelligence:</b> Diseñé dashboards interactivos y modelos de datos en Power BI y Tableau para soporte en toma de decisiones.", styles['CVBullet']))

    # Freelance / Proyectos
    t_data3 = [
        [Paragraph("<b>Data Engineer</b> — Proyectos de Automatización & Cloud (Freelance)", styles['CVJobTitle']),
         Paragraph("2024 – Presente", styles['CVJobMeta'])]
    ]
    t3 = Table(t_data3, colWidths=[395, 160])
    t3.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 0), ('TOPPADDING', (0,0), (-1,-1), 0), ('LEFTPADDING', (0,0), (-1,-1), 0)]))
    story.append(t3)
    story.append(Paragraph("• <b>Pipelines IA (n8n & LLMs):</b> Flujos en n8n integrando APIs (Gemini) para procesar, limpiar y calificar leads B2B en CRM (+40% velocidad de ciclo).", styles['CVBullet']))
    story.append(Paragraph("• <b>Infraestructura Cloud:</b> Despliegue de bases de datos resilientes y alta disponibilidad en AWS, GCP y Azure en producción (jofastsa.com).", styles['CVBullet']))
    story.append(Paragraph("• <b>Big Data:</b> Procesamiento distribuido y pruebas de arquitecturas escalables con Apache Spark, Hadoop y Docker.", styles['CVBullet']))

    # 3. EDUCACIÓN
    story.append(Paragraph("EDUCACIÓN", styles['CVSectionHeader']))
    story.append(HRFlowable(width="100%", thickness=0.4, color=c_line, spaceBefore=1, spaceAfter=2))
    
    t_edu1 = [
        [Paragraph("<b>Ing. de Sistemas Computacionales</b> — Universidad Privada del Norte (UPN)", styles['CVJobTitle']),
         Paragraph("2026 – Presente (En curso)", styles['CVJobMeta'])]
    ]
    te1 = Table(t_edu1, colWidths=[420, 135])
    te1.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 0), ('TOPPADDING', (0,0), (-1,-1), 0), ('LEFTPADDING', (0,0), (-1,-1), 0)]))
    story.append(te1)
    
    t_edu2 = [
        [Paragraph("<b>Técnico en Desarrollo de Sistemas de Información</b> — Inst. Sabio Nacional Antúnez de Mayolo", styles['CVJobTitle']),
         Paragraph("2022 – 2025 (Egresado)", styles['CVJobMeta'])]
    ]
    te2 = Table(t_edu2, colWidths=[420, 135])
    te2.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('BOTTOMPADDING', (0,0), (-1,-1), 0), ('TOPPADDING', (0,0), (-1,-1), 0), ('LEFTPADDING', (0,0), (-1,-1), 0)]))
    story.append(te2)

    # 4. CERTIFICACIONES Y ESPECIALIZACIONES
    story.append(Paragraph("CERTIFICACIONES Y ESPECIALIZACIONES", styles['CVSectionHeader']))
    story.append(HRFlowable(width="100%", thickness=0.4, color=c_line, spaceBefore=1, spaceAfter=2))
    story.append(Paragraph("• <b>SAP MM - Logística y Materiales (Key User / Usuario Experto):</b> Global Tecnologías Academy (2026).", styles['CVBullet']))
    story.append(Paragraph("• <b>Ciencia de Datos aplicada a la Logística (Specialty):</b> Ingenium (2025). Modelos predictivos y optimización de suministros.", styles['CVBullet']))
    story.append(Paragraph("• <b>Data Science y Machine Learning (Specialty):</b> Toulouse Lautrec (2024 – 2025). MLOps, Scikit-Learn, Deep Learning.", styles['CVBullet']))

    # 5. HABILIDADES TÉCNICAS
    story.append(Paragraph("HABILIDADES TÉCNICAS & HERRAMIENTAS", styles['CVSectionHeader']))
    story.append(HRFlowable(width="100%", thickness=0.4, color=c_line, spaceBefore=1, spaceAfter=2))
    
    skills_data = [
        [Paragraph("<b>Desarrollo & APIs:</b>", styles['CVSkillLabel']), Paragraph("PHP, Angular, API REST, Microservicios, Java, FastAPI, JavaScript, HTML5/CSS3.", styles['CVSkillVal'])],
        [Paragraph("<b>Data Engineering:</b>", styles['CVSkillLabel']), Paragraph("Data Pipelines (ETL/ELT), Python, SQL, Apache Spark, Hadoop, n8n, Make, Web Scraping.", styles['CVSkillVal'])],
        [Paragraph("<b>Bases de Datos:</b>", styles['CVSkillLabel']), Paragraph("PostgreSQL, MySQL, MongoDB, Data Warehousing, Modelado Relacional y Dimensional.", styles['CVSkillVal'])],
        [Paragraph("<b>Cloud & DevOps:</b>", styles['CVSkillLabel']), Paragraph("AWS, Google Cloud Platform (GCP), Microsoft Azure, Docker, Git, GitHub Actions, CI/CD.", styles['CVSkillVal'])],
        [Paragraph("<b>BI & Visualización:</b>", styles['CVSkillLabel']), Paragraph("Power BI (DAX, Power Query), Tableau, Streamlit, Metabase.", styles['CVSkillVal'])],
        [Paragraph("<b>ERP & Negocio:</b>", styles['CVSkillLabel']), Paragraph("SAP MM (Key User), Gestión de Materiales, Control de Inventarios, Optimización Logística.", styles['CVSkillVal'])],
    ]
    t_skills = Table(skills_data, colWidths=[110, 445])
    t_skills.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 1),
        ('TOPPADDING', (0,0), (-1,-1), 0.5),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(t_skills)
    
    doc.build(story)
    print("public/cv.pdf generated successfully.")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
